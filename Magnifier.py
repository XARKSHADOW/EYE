# # #Magnifier.py
# import tkinter as tk
# from tkinter import filedialog
# from PIL import Image, ImageTk, ImageDraw
#
#
# class MagnifierApp:
#     def __init__(self, roots, ):
#         self.upload_file = ""
#         self.root = roots
#         self.root.title("Magnifier App")
#
#         # Your other initializations
#         self.upload_button = tk.Button(self.root, text="Upload File", command=self.upload_file)
#         self.upload_button.place(x=10, y=10)
#
#         def upload_file(self):
#             filename = filedialog.askopenfilename(
#                 filetypes=[("All Files", "*.*"), ("Image Files", "*.png;*.jpg;*.jpeg;*.bmp;*.gif"),
#                            ("PDF Files", "*.pdf"),
#                            ("Word Documents", "*.docx")])
#             if filename:
#                 file_ext = filename.split('.')[-1].lower()
#                 if file_ext in ['png', 'jpg', 'jpeg', 'bmp', 'gif']:
#                     self.img = Image.open(filename)
#                     self.tk_img = ImageTk.PhotoImage(self.img)
#                     self.canvas.create_image(self.canvas.winfo_width() // 2, self.canvas.winfo_height() // 2,
#                                              anchor=tk.CENTER, image=self.tk_img)
#
#                 elif file_ext == 'pdf':
#                     import PyPDF2
#                     pdf_file = open(filename, 'rb')
#                     pdf_reader = PyPDF2.PdfFileReader(pdf_file)
#                     page = pdf_reader.getPage(0)
#                     text = page.extract_text()
#                     img = Image.new('RGB', (800, 1000), color=(255, 255, 255))
#                     d = ImageDraw.Draw(img)
#                     d.text((10, 10), text, fill=(0, 0, 0))
#                     self.tk_img = ImageTk.PhotoImage(img)
#                     self.canvas.create_image(0, 0, anchor=tk.NW, image=self.tk_img)
#                 elif file_ext == 'docx':
#                     import docx
#                     doc = docx.Document(filename)
#                     full_text = []
#                     for para in doc.paragraphs:
#                         full_text.append(para.text)
#                     text = '\n'.join(full_text)
#                     img = Image.new('RGB', (800, 1000), color=(255, 255, 255))
#                     d = ImageDraw.Draw(img)
#                     d.text((10, 10), text, fill=(0, 0, 0))
#                     self.tk_img = ImageTk.PhotoImage(img)
#                     self.canvas.create_image(0, 0, anchor=tk.NW, image=self.tk_img)
#
#
#
#         def upload():
#           fileName = upload_file(self)
#
#         # Your upload_file code here
#
#         # Get screen dimensions
#         screen_width = self.root.winfo_screenwidth()
#         screen_height = self.root.winfo_screenheight()
#
#         # Calculate a suitable window size
#         window_width = max(800, screen_width - 100)
#         window_height = max(600, screen_height - 100)
#
#         # Center the window
#         x_position = (screen_width - window_width) // 2
#         y_position = (screen_height - window_height) // 2
#
#         self.root.geometry(f"{window_width}x{window_height}+{x_position}+{y_position}")
#
#         self.canvas = tk.Canvas(self.root, width=window_width, height=window_height, bg="white")
#         self.canvas.pack(fill=tk.BOTH, expand=True)
#
#         self.magnifier_size = 200
#         self.magnifier_zoom = 2
#
#
#         self.img = Image.open("/home/zephyrdarkfire/PycharmProjects/Zephyr/Tracker/examples_image.jpg")
#         self.tk_img = ImageTk.PhotoImage(self.img)
#         self.canvas.create_image(0, 0, anchor=tk.NW, image=self.tk_img)
#
#         self.canvas.bind("<Motion>", self.magnify)
#
#         # Add upload button
#         # self.upload_button = tk.Button(root, text="Upload File", command=self.upload_file)
#         # self.upload_button.place(x=10, y=10)
#
#         # Add a panel for the upload button
#         self.panel = tk.Frame(root, bg="green")
#         self.panel.pack(side=tk.TOP, fill=tk.X)
#
#         # Add upload button to the panel
#         menu = tk.Menu(root)
#         root.config(menu=menu)
#         options_menu = tk.Menu(menu, tearoff=0)
#         menu.add_cascade(label="Options", menu=options_menu)
#         options_menu.add_command(label="Upload file", command=upload)
#         self.upload_button = tk.Button(self.panel, text="Upload File", command=self.upload_file)
#         self.upload_button.pack(padx=10, pady=10)
#
#         # Upload Feature
#
#     def magnify(self, event):
#         x, y = event.x, event.y
#         box = (x - self.magnifier_size // 2, y - self.magnifier_size // 2, x + self.magnifier_size // 2,
#                y + self.magnifier_size // 2)
#         region = self.img.crop(box)
#         region = region.resize((self.magnifier_size * self.magnifier_zoom, self.magnifier_size * self.magnifier_zoom))
#         # Create a circular mask
#         mask = Image.new('L', region.size, 0)
#         draw = ImageDraw.Draw(mask)
#         draw.ellipse((0, 0) + region.size, fill=600)
#
#         # Apply the mask to the region
#         region.putalpha(mask)
#         magnified_img = ImageTk.PhotoImage(region)
#         self.canvas.create_image(x, y, image=magnified_img)
#         self.canvas.image = magnified_img
#
#
# if __name__ == "__main__":
#     root = tk.Tk()
#     app = MagnifierApp(root)
#     root.mainloop()
# The working code here
import tkinter as tk
from tkinter import filedialog
from PIL import Image, ImageTk, ImageDraw
import PyPDF2
import docx


class MagnifierApp:
    def __init__(self, root):
        self.hide_magnifier = None
        self.root = root
        self.root.title("Magnifier App")

        # Your other initializations
        self.upload_button = tk.Button(self.root, text="Upload File", command=self.upload_file)
        self.upload_button.place(x=10, y=10)

        # Get screen dimensions
        screen_width = self.root.winfo_screenwidth()
        screen_height = self.root.winfo_screenheight()

        # Calculate a suitable window size
        window_width = max(800, screen_width - 100)
        window_height = max(600, screen_height - 100)

        # Center the window
        x_position = (screen_width - window_width) // 2
        y_position = (screen_height - window_height) // 2

        self.root.geometry(f"{window_width}x{window_height}+{x_position}+{y_position}")

        self.canvas = tk.Canvas(self.root, width=window_width, height=window_height, bg="black")
        self.canvas.pack(fill=tk.BOTH, expand=True)

        self.magnifier_size = 200
        self.magnifier_zoom = 2

        # Initialize content placeholders
        self.img = None
        self.tk_img = None

        self.canvas.bind("<Motion>", self.magnify)
        self.canvas.bind("<Leave>", self.hide_magnifier)

        # Add a panel for the upload button
        self.panel = tk.Frame(root, bg="green")
        self.panel.pack(side=tk.TOP, fill=tk.X)

        # Add upload button to the panel
        menu = tk.Menu(root)
        root.config(menu=menu)
        options_menu = tk.Menu(menu, tearoff=0)
        menu.add_cascade(label="Options", menu=options_menu)
        options_menu.add_command(label="Upload file", command=self.upload_file)

    def upload_file(self):
        filename = filedialog.askopenfilename(
            filetypes=[("All Files", "*.*"), ("Image Files", "*.png;*.jpg;*.jpeg;*.bmp;*.gif"),
                       ("PDF Files", "*.pdf"), ("Word Documents", "*.docx")])
        if filename:
            file_ext = filename.split('.')[-1].lower()
            if file_ext in ['png', 'jpg', 'jpeg', 'bmp', 'gif']:
                self.img = Image.open(filename)
                self.render_image_on_canvas()
            elif file_ext == 'pdf':
                self.img = self.render_pdf_to_image(filename)
                self.render_image_on_canvas()
            elif file_ext == 'docx':
                self.img = self.render_docx_to_image(filename)
                self.render_image_on_canvas()

    def render_image_on_canvas(self):
        # Calculate dimensions for fitting image to canvas
        img_width, img_height = self.img.size
        canvas_width = self.canvas.winfo_width()
        canvas_height = self.canvas.winfo_height()

        # Calculate resizing factor to fit the image within the canvas
        resize_factor = min(canvas_width // img_width, canvas_height // img_height)
        new_width = int(img_width * resize_factor)
        new_height = int(img_height * resize_factor)

        # Resize image
        resized_img = self.img.resize((new_width, new_height))

        # Convert to ImageTk format
        self.tk_img = ImageTk.PhotoImage(resized_img)

        # Clear existing canvas items
        self.canvas.delete("all")

        # Center image on canvas
        x = (canvas_width - new_width) // 2
        y = (canvas_height - new_height) // 2

        # Display image on canvas
        self.canvas.create_image(x, y, anchor=tk.NW, image=self.tk_img)

    def render_pdf_to_image(self, filename):
        pdf_file = open(filename, 'rb')
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        page = pdf_reader.pages[0]
        text = page.extract_text()
        img = Image.new('RGB', (800, 1000), color=(255, 255, 255))
        d = ImageDraw.Draw(img)
        d.text((10, 10), text, fill=(0, 0, 0))
        return img

    def render_docx_to_image(self, filename):
        doc = docx.Document(filename)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        text = '\n'.join(full_text)
        img = Image.new('RGB', (800, 1000), color=(255, 255, 255))
        d = ImageDraw.Draw(img)
        d.text((10, 10), text, fill=(0, 0, 0))
        return img

    def magnify(self, event):
        if not self.img:
            return

        img_width, img_height = self.img.size
        x, y = event.x, event.y

        # Calculate the magnification box based on mouse position
        magnifier_half_size = self.magnifier_size // 2
        box_left = max(0, x - magnifier_half_size)
        box_top = max(0, y - magnifier_half_size)
        box_right = min(img_width, x + magnifier_half_size)
        box_bottom = min(img_height, y + magnifier_half_size)

        # Adjust box if it goes out of bounds
        if box_right - box_left < self.magnifier_size:
            if box_right == img_width:
                box_left = img_width - self.magnifier_size
            else:
                box_right = box_left + self.magnifier_size

        if box_bottom - box_top < self.magnifier_size:
            if box_bottom == img_height:
                box_top = img_height - self.magnifier_size
            else:
                box_bottom = box_top + self.magnifier_size

        # Create a blank white image to cover the entire canvas
        cover_img = Image.new('RGB', (self.canvas.winfo_width(), self.canvas.winfo_height()), color=(255, 255, 255))
        cover_tk_img = ImageTk.PhotoImage(cover_img)

        # Display the covered image on the canvas
        self.canvas.create_image(0, 0, anchor=tk.NW, image=cover_tk_img)
        self.canvas.image = cover_tk_img

        # Apply the magnifier effect on the covered image
        region = self.img.crop((box_left, box_top, box_right, box_bottom))
        region = region.resize((self.magnifier_size * self.magnifier_zoom, self.magnifier_size * self.magnifier_zoom))

        # Create a circular mask
        mask = Image.new('L', region.size, 0)
        draw = ImageDraw.Draw(mask)
        draw.ellipse((0, 0) + region.size, fill=255)

        # Apply the mask to the region
        region.putalpha(mask)

        # Convert to PhotoImage
        magnified_img = ImageTk.PhotoImage(region)

        # Display the magnified image on the canvas
        self.canvas.create_image(x, y, image=magnified_img)
        self.canvas.image = magnified_img


if __name__ == "__main__":
    root = tk.Tk()
    app = MagnifierApp(root)
    root.mainloop()

